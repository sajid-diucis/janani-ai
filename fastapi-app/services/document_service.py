"""
Document Processing Service
Handles Word document uploads for medical history and patient profile
"""
import os
import json
import tempfile
import re
from typing import Dict, List, Optional
from datetime import datetime

try:
    from docx import Document
except ImportError:
    Document = None

from services.ai_service import AIService
from models.care_models import DocumentProfile, ValidationStatus


class DocumentService:
    def __init__(self):
        self.ai_service = AIService()
        # Store profiles with persistence
        self.db_path = os.path.join("data", "document_profiles.json")
        os.makedirs("data", exist_ok=True)
        self.user_profiles: Dict[str, dict] = self._load_profiles()

    def _load_profiles(self) -> dict:
        """Load profiles from JSON file"""
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                print(f"Error loading document profiles: {e}")
        return {}

    def _save_profiles(self):
        """Save profiles to JSON file"""
        try:
            with open(self.db_path, "w", encoding="utf-8") as f:
                json.dump(self.user_profiles, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error saving document profiles: {e}")
    
    def _extract_text_from_doc(self, file_content: bytes) -> str:
        """
        Extract text from old .doc format (binary)
        Uses basic text extraction from binary data
        """
        try:
            # Try to find text in the binary .doc file
            # .doc files store text in specific locations
            text_parts = []
            
            # Decode with latin-1 to preserve all bytes, then extract readable text
            content_str = file_content.decode('latin-1', errors='ignore')
            
            # Find readable text patterns (words with Bengali or English chars)
            # Look for sequences of printable characters
            readable_chunks = re.findall(r'[\x20-\x7E\u0980-\u09FF]{4,}', content_str)
            
            # Filter out binary garbage and keep meaningful text
            filtered_text = []
            for chunk in readable_chunks:
                # Skip chunks that look like file paths or binary artifacts
                if '\\' not in chunk or len(chunk) > 50:
                    # Clean up the chunk
                    clean = chunk.strip()
                    if clean and len(clean) > 3:
                        filtered_text.append(clean)
            
            return ' '.join(filtered_text)
        except Exception as e:
            return ""
    
    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from .docx format using python-docx"""
        if Document is None:
            return ""
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            os.unlink(tmp_path)
            return "\n".join(full_text)
        except Exception as e:
            try:
                os.unlink(tmp_path)
            except:
                pass
            return ""
    
    async def process_document(self, file_content: bytes, filename: str, user_id: str = "default_user") -> dict:
        """
        Process uploaded Word document and extract medical/financial information
        """
        try:
            # Determine file type by extension
            is_docx = filename.lower().endswith('.docx')
            is_doc = filename.lower().endswith('.doc')
            
            document_text = ""
            
            if is_docx:
                # Try .docx extraction
                document_text = self._extract_text_from_docx(file_content)
                if not document_text:
                    # Fallback to binary extraction
                    document_text = self._extract_text_from_doc(file_content)
            elif is_doc:
                # Old .doc format - use binary extraction
                document_text = self._extract_text_from_doc(file_content)
            else:
                return {
                    "success": False,
                    "error": "Unsupported file format",
                    "message_bengali": "এই ফরম্যাট সাপোর্ট করে না। .doc বা .docx ফাইল দিন।"
                }
            
            if not document_text or len(document_text.strip()) < 10:
                return {
                    "success": False,
                    "error": "Could not extract text from document",
                    "message_bengali": "ডকুমেন্ট থেকে তথ্য বের করা যাচ্ছে না। অনুগ্রহ করে .docx ফরম্যাটে সেভ করে আবার চেষ্টা করুন।"
                }
            
            # Use AI to extract structured information
            extracted_info = await self._extract_info_with_ai(document_text)
            
            # Generate Bengali summary
            bengali_summary = self._generate_bengali_summary(extracted_info)
            
            # Store the profile using the new P0 model (Refined Type Safety)
            doc_profile = DocumentProfile(
                filename=filename,
                raw_text_snippet=document_text[:2000],
                structured_info=extracted_info,
                extraction_confidence=0.85, # In production, pull from LLM metadata
                validation_status=ValidationStatus.PENDING,
                extracted_by_model="gemini-1.5-flash"
            )
            
            # Serialize for JSON persistence
            try:
                dumped = doc_profile.model_dump(mode='json')
            except AttributeError:
                dumped = json.loads(doc_profile.json())
                
            # APPEND to user's document history (P0 Multi-Document Support)
            if user_id not in self.user_profiles:
                self.user_profiles[user_id] = []
            
            # If the user has old-style single profile, convert it
            if isinstance(self.user_profiles[user_id], dict):
                self.user_profiles[user_id] = [self.user_profiles[user_id]]
                
            self.user_profiles[user_id].append(dumped)
            
            # Persist to disk
            self._save_profiles()
            
            return {
                "success": True,
                "message_bengali": "আপনার তথ্য সফলভাবে আপলোড হয়েছে!",
                "extracted_info": extracted_info,
                "bengali_summary": bengali_summary
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message_bengali": f"ডকুমেন্ট প্রসেস করতে সমস্যা হয়েছে: {str(e)}"
            }
    
    async def _extract_info_with_ai(self, document_text: str) -> dict:
        """
        Use AI to extract structured information from document text
        """
        prompt = f"""Analyze this medical/health document of a pregnant woman and extract information in JSON format.

Document text:
{document_text[:3000]}

Extract and return ONLY valid JSON (no other text):
{{
    "patient_name": "name if found or null",
    "age": "age as number or null",
    "trimester": "first/second/third or null",
    "week": "pregnancy week as number or null",
    "blood_group": "blood group or null",
    "weight_kg": "weight in kg as number or null",
    "height_cm": "height in cm as number or null",
    
    "medical_conditions": ["list of conditions like anemia, diabetes, hypertension, thyroid, etc"],
    "past_medical_history": ["list of past illnesses or surgeries"],
    "family_history": ["list of family health conditions"],
    "allergies": ["list of food or medicine allergies"],
    "current_medications": ["list of current medicines"],
    
    "hemoglobin_level": "Hb level as number or null",
    "blood_sugar_fasting": "fasting sugar level or null",
    "blood_pressure": "BP reading or null",
    
    "monthly_income_bdt": "income in BDT as number or null",
    "food_budget_weekly_bdt": "weekly food budget or null",
    
    "dietary_preferences": ["vegetarian, non-veg, halal, etc"],
    "foods_disliked": ["foods the patient doesn't like"],
    
    "previous_pregnancies": "number or null",
    "complications_history": ["any past pregnancy complications"],
    
    "notes": "any other important information"
}}

If information is not found, use null. Extract whatever is available."""

        try:
            response = await self.ai_service.get_response(
                message=prompt,
                conversation_history=[],
                is_emergency=False,
                user_context=None
            )
            
            # Parse JSON from response
            import re
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                return json.loads(json_match.group())
            else:
                return {"raw_extraction": response, "parse_error": True}
                
        except Exception as e:
            return {"extraction_error": str(e)}
    
    def _generate_bengali_summary(self, info: dict) -> str:
        """
        Generate a Bengali summary of extracted information
        """
        summary_parts = []
        
        if info.get("patient_name"):
            summary_parts.append(f"নাম: {info['patient_name']}")
        
        if info.get("age"):
            summary_parts.append(f"বয়স: {info['age']} বছর")
        
        if info.get("trimester"):
            trimester_bn = {"first": "১ম", "second": "২য়", "third": "৩য়"}.get(info['trimester'], info['trimester'])
            summary_parts.append(f"গর্ভাবস্থা: {trimester_bn} ত্রৈমাসিক")
        
        if info.get("medical_conditions"):
            conditions = info['medical_conditions']
            if conditions:
                cond_map = {
                    "anemia": "রক্তস্বল্পতা",
                    "diabetes": "ডায়াবেটিস",
                    "hypertension": "উচ্চ রক্তচাপ",
                    "thyroid": "থাইরয়েড"
                }
                conds_bn = [cond_map.get(c.lower(), c) for c in conditions]
                summary_parts.append(f"স্বাস্থ্য সমস্যা: {', '.join(conds_bn)}")
        
        if info.get("allergies"):
            summary_parts.append(f"এলার্জি: {', '.join(info['allergies'])}")
        
        if info.get("hemoglobin_level"):
            summary_parts.append(f"হিমোগ্লোবিন: {info['hemoglobin_level']} g/dL")
        
        if info.get("food_budget_weekly_bdt"):
            summary_parts.append(f"সাপ্তাহিক খাবার বাজেট: ৳{info['food_budget_weekly_bdt']}")
        
        if not summary_parts:
            return "তথ্য পাওয়া গেছে। আপনার প্রোফাইল আপডেট হয়েছে।"
        
        return " | ".join(summary_parts)
    
    def get_user_profile(self, user_id: str = "default_user") -> Optional[dict]:
        """
        Get stored user profile
        """
        return self.user_profiles.get(user_id)
    
    def get_conditions_from_profile(self, user_id: str = "default_user") -> List[str]:
        """
        Get medical conditions from stored profile
        """
        profile = self.user_profiles.get(user_id)
        if profile and profile.get("structured_info"):
            return profile["structured_info"].get("medical_conditions", [])
        return []
    
    def get_allergies_from_profile(self, user_id: str = "default_user") -> List[str]:
        """
        Get allergies from stored profile
        """
        profile = self.user_profiles.get(user_id)
        if profile and profile.get("structured_info"):
            return profile["structured_info"].get("allergies", [])
        return []
    
    def get_budget_from_profile(self, user_id: str = "default_user") -> Optional[float]:
        """
        Get food budget from stored profile (checks newest documents first)
        """
        docs = self.user_profiles.get(user_id, [])
        if not docs:
            return None
            
        if isinstance(docs, dict): docs = [docs]
        
        # Search newest first
        for doc in reversed(docs):
            budget = doc.get("structured_info", {}).get("food_budget_weekly_bdt")
            if budget: return budget
        return None

    def get_combined_profile(self, user_id: str, existing_profile_dict: Optional[dict] = None) -> dict:
        """
        Merge extracted data from MANY documents with an existing profile dictionary.
        This provides a cumulative "Living Context" for RAG.
        """
        docs = self.user_profiles.get(user_id, [])
        if not docs:
            return existing_profile_dict or {}
            
        if isinstance(docs, dict): docs = [docs] # Compatibility
        
        combined = (existing_profile_dict or {}).copy()
        
        key_map = {
            "patient_name": "name",
            "age": "age",
            "blood_group": "blood_group",
            "weight_kg": "current_weight_kg",
            "height_cm": "height_cm",
            "hemoglobin_level": "hemoglobin_level",
            "blood_pressure": "blood_pressure_systolic", # Simple BP mapping
            "monthly_income_bdt": "monthly_income_bdt",
            "week": "current_week"
        }
        
        # Iterate documents (Oldest to Newest) so newest data wins conflicts
        for doc in docs:
            extracted = doc.get("structured_info", {})
            for doc_key, profile_key in key_map.items():
                val = extracted.get(doc_key)
                # Update if current value is default/missing or if it's a newer doc
                if val is not None:
                    combined[profile_key] = val
            
            # Accumulate lists (Medical Conditions, Allergies, etc.)
            def merge_lists(key_p, key_d):
                p_list = combined.get(key_p, [])
                d_list = extracted.get(key_d, [])
                # Combine and unique
                combined[key_p] = list(set(p_list + d_list))

            merge_lists("existing_conditions", "medical_conditions")
            merge_lists("allergies", "allergies")
            merge_lists("current_medications", "current_medications")
            merge_lists("previous_complications", "complications_history")
            
            # Detect and set fast-track flags from ANY doc
            doc_text_lower = str(extracted.get("medical_conditions", [])).lower()
            if any(x in doc_text_lower for x in ["diabetes", "gdm", "ডায়াবেটিস"]):
                combined["has_gestational_diabetes"] = True
            if any(x in doc_text_lower for x in ["hypertension", "high bp", "উচ্চ রক্তচাপ"]):
                combined["has_hypertension"] = True
            if any(x in doc_text_lower for x in ["anemia", "অ্যানিমিয়া", "রক্তাল্পতা"]):
                combined["has_anemia"] = True
            
            if extracted.get("emergency_contact"):
                combined["emergency_contact_phone"] = extracted.get("emergency_contact")

        # Set metadata from last doc
        last_doc = docs[-1]
        combined["_is_augmented_by_doc"] = True
        combined["_doc_filename"] = last_doc.get("filename")
        combined["_all_docs_count"] = len(docs)
        
        return combined


# Global instance
document_service = DocumentService()
